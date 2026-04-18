from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template



class TemplateService:
    """Word文档模板渲染服务"""

    def __init__(self, template_dir: str = "templates/reports"):
        self.template_dir = Path(template_dir)
        self.template_dir.mkdir(parents=True, exist_ok=True)

    def get_chapter_template(self, report_type: str) -> Dict[str, str]:
        """获取报告章节模板"""
        templates = {
            "feasibility": {
                "1": "## {{ chapter_num }}. 项目概述\n\n### {{ chapter }}.1 项目背景\n{{ content }}",
                "2": "## {{ chapter_num }}. 工程建设的必要性\n\n### {{ chapter }}.1 项目由来\n{{ content }}",
                "3": "## {{ chapter_num }}. 工程任务与规模\n\n### {{ chapter }}.1 防洪标准\n{{ content }}",
            }
        }
        return templates.get(report_type, {})

    def render_chapter(
        self,
        template: str,
        chapter_num: str,
        content: str,
        context: Dict[str, Any]
    ) -> str:
        """渲染单个章节"""
        t = Template(template, autoescape=True)
        return t.render(
            chapter_num=chapter_num,
            chapter=chapter_num,
            content=content,
            **context
        )

    def create_word_document(
        self,
        chapters: Dict[str, str],
        output_path: str,
        title: str
    ) -> str:
        """创建Word文档"""
        doc = Document()

        # 设置标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加章节
        for chapter_title, chapter_content in chapters.items():
            doc.add_heading(chapter_title, level=1)
            doc.add_paragraph(chapter_content)

        # 保存
        doc.save(output_path)
        return output_path

    def template_exists(self, report_type: str) -> bool:
        """检查模板是否存在"""
        template_path = self.template_dir / f"{report_type}_template.md"
        return template_path.exists()